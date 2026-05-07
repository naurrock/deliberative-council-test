"""
Output formatting for Deliberative Council.

Converts FinalReport to various output formats:
- Markdown (internal representation)
- PDF (via weasyprint)
- DOCX (via python-docx)
- JSON (structured)
- Plain text
"""

from __future__ import annotations

import json
import logging
from pathlib import Path

from council.types import FinalReport

logger = logging.getLogger(__name__)


def format_report(report: FinalReport, fmt: str = "markdown") -> str:
    """Format a FinalReport into the specified output format.

    Args:
        report: The final report to format.
        fmt: Output format — "markdown", "pdf", "docx", "json", "text".

    Returns:
        Formatted string (for text formats) or file path (for binary formats).
    """
    if fmt == "markdown":
        return report.raw_markdown or _to_markdown(report)
    elif fmt == "json":
        return _to_json(report)
    elif fmt == "text":
        return _to_text(report)
    elif fmt == "pdf":
        return _to_pdf(report)
    elif fmt == "docx":
        return _to_docx(report)
    else:
        raise ValueError(f"Unknown output format: {fmt}")


def save_report(report: FinalReport, output_path: str | Path, fmt: str | None = None) -> str:
    """Save a FinalReport to a file.

    Args:
        report: The final report to save.
        output_path: File path to save to.
        fmt: Output format. If None, inferred from file extension.

    Returns:
        The path the file was saved to.
    """
    path = Path(output_path)

    if fmt is None:
        ext_map = {
            ".md": "markdown",
            ".pdf": "pdf",
            ".docx": "docx",
            ".json": "json",
            ".txt": "text",
        }
        fmt = ext_map.get(path.suffix.lower(), "markdown")

    content = format_report(report, fmt)

    # For binary formats, format_report returns the file path
    if fmt in ("pdf", "docx"):
        # Content is the path to the generated file
        if content != str(path):
            import shutil
            shutil.copy2(content, path)
        return str(path)

    # For text formats, write content directly
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    return str(path)


# ── Format Implementations ───────────────────────────────────────────────


def _to_markdown(report: FinalReport) -> str:
    """Convert report to Markdown."""
    lines = []
    lines.append(f"# {report.question}")
    lines.append("")
    lines.append(f"**Complexity**: {report.complexity.value}  ")
    lines.append(f"**Rounds**: {report.rounds_completed}  ")
    lines.append(f"**Convergence**: {report.convergence_score:.2f}  ")
    lines.append("")

    lines.append(report.answer)
    lines.append("")

    if report.key_points:
        lines.append("## Key Points")
        lines.append("")
        for kp in report.key_points:
            consensus_marker = {
                "strong": "[STRONG]",
                "moderate": "[MODERATE]",
                "contested": "[CONTESTED]",
            }.get(kp.consensus.value, "[?]")
            lines.append(f"- {consensus_marker} {kp.point}")
            if kp.dissent:
                lines.append(f"  - Dissent: {kp.dissent}")
        lines.append("")

    if report.dissenting_views:
        lines.append("## Dissenting Views")
        lines.append("")
        for dv in report.dissenting_views:
            lines.append(f"- {dv}")
        lines.append("")

    if report.research_sources:
        lines.append("## Sources")
        lines.append("")
        for src in report.research_sources:
            lines.append(f"- {src}")
        lines.append("")

    if report.futility_notes:
        lines.append("## Notes")
        lines.append("")
        lines.append(report.futility_notes)
        lines.append("")

    lines.append("## Pipeline Trace")
    lines.append("")
    trace = report.pipeline_trace
    lines.append(f"| Phase | Tokens |")
    lines.append(f"|-------|--------|")
    lines.append(f"| Scout | {trace.scout_tokens:,} |")
    lines.append(f"| Research | {trace.research_tokens:,} |")
    lines.append(f"| Debate | {trace.debate_tokens:,} |")
    lines.append(f"| Synthesis | {trace.synthesis_tokens:,} |")
    lines.append(f"| **Total** | **{trace.total_tokens:,}** |")
    lines.append("")

    if trace.models_used:
        lines.append("### Models Used")
        lines.append("")
        for role, mu in trace.models_used.items():
            lines.append(f"- **{role}**: {mu.model} ({mu.family}) — {mu.tokens:,} tokens")

    return "\n".join(lines)


def _to_json(report: FinalReport) -> str:
    """Convert report to JSON."""
    return report.model_dump_json(indent=2)


def _to_text(report: FinalReport) -> str:
    """Convert report to plain text (no markdown)."""
    lines = []
    lines.append(report.question)
    lines.append("=" * len(report.question))
    lines.append("")
    lines.append(f"Complexity: {report.complexity.value}")
    lines.append(f"Rounds: {report.rounds_completed}")
    lines.append(f"Convergence: {report.convergence_score:.2f}")
    lines.append("")

    # Strip markdown from answer for plain text
    import re
    answer_text = report.answer
    answer_text = re.sub(r'#+ ', '', answer_text)
    answer_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', answer_text)
    answer_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', answer_text)
    lines.append(answer_text)
    lines.append("")

    if report.key_points:
        lines.append("KEY POINTS")
        lines.append("-" * 10)
        for kp in report.key_points:
            lines.append(f"  [{kp.consensus.value.upper()}] {kp.point}")
        lines.append("")

    if report.dissenting_views:
        lines.append("DISSENTING VIEWS")
        lines.append("-" * 15)
        for dv in report.dissenting_views:
            lines.append(f"  - {dv}")
        lines.append("")

    if report.research_sources:
        lines.append("SOURCES")
        lines.append("-" * 7)
        for src in report.research_sources:
            lines.append(f"  - {src}")
        lines.append("")

    return "\n".join(lines)


def _to_pdf(report: FinalReport) -> str:
    """Convert report to PDF using weasyprint."""
    import tempfile

    md_content = _to_markdown(report)
    html_content = _markdown_to_html(md_content)

    output_path = tempfile.mktemp(suffix=".pdf")
    try:
        from weasyprint import HTML

        HTML(string=html_content).write_pdf(output_path)
        return output_path
    except ImportError:
        logger.warning("weasyprint not available, falling back to markdown")
        # Save as markdown instead
        fallback_path = output_path.replace(".pdf", ".md")
        Path(fallback_path).write_text(md_content, encoding="utf-8")
        return fallback_path
    except Exception as e:
        logger.error(f"PDF generation failed: {e}")
        fallback_path = output_path.replace(".pdf", ".md")
        Path(fallback_path).write_text(md_content, encoding="utf-8")
        return fallback_path


def _to_docx(report: FinalReport) -> str:
    """Convert report to DOCX using python-docx."""
    import tempfile

    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(report.question, level=1)

    # Metadata
    p = doc.add_paragraph()
    p.add_run(f"Complexity: {report.complexity.value}  |  ")
    p.add_run(f"Rounds: {report.rounds_completed}  |  ")
    p.add_run(f"Convergence: {report.convergence_score:.2f}")

    # Main answer
    doc.add_heading("Answer", level=2)
    # Split answer by paragraphs and add each
    for para in report.answer.split("\n\n"):
        if para.strip().startswith("#"):
            level = para.strip().count("#")
            text = para.strip().lstrip("#").strip()
            doc.add_heading(text, level=min(level, 4))
        elif para.strip().startswith("- "):
            for item in para.strip().split("\n"):
                item = item.lstrip("- ").strip()
                if item:
                    doc.add_paragraph(item, style="List Bullet")
        else:
            doc.add_paragraph(para.strip())

    # Key points
    if report.key_points:
        doc.add_heading("Key Points", level=2)
        for kp in report.key_points:
            text = f"[{kp.consensus.value.upper()}] {kp.point}"
            doc.add_paragraph(text, style="List Bullet")
            if kp.dissent:
                doc.add_paragraph(f"Dissent: {kp.dissent}", style="List Bullet 2")

    # Dissenting views
    if report.dissenting_views:
        doc.add_heading("Dissenting Views", level=2)
        for dv in report.dissenting_views:
            doc.add_paragraph(dv, style="List Bullet")

    # Sources
    if report.research_sources:
        doc.add_heading("Sources", level=2)
        for src in report.research_sources:
            doc.add_paragraph(src, style="List Bullet")

    # Pipeline trace
    doc.add_heading("Pipeline Trace", level=2)
    trace = report.pipeline_trace
    table = doc.add_table(rows=5, cols=2)
    table.rows[0].cells[0].text = "Phase"
    table.rows[0].cells[1].text = "Tokens"
    table.rows[1].cells[0].text = "Scout"
    table.rows[1].cells[1].text = f"{trace.scout_tokens:,}"
    table.rows[2].cells[0].text = "Research"
    table.rows[2].cells[1].text = f"{trace.research_tokens:,}"
    table.rows[3].cells[0].text = "Debate"
    table.rows[3].cells[1].text = f"{trace.debate_tokens:,}"
    table.rows[4].cells[0].text = "Synthesis"
    table.rows[4].cells[1].text = f"{trace.synthesis_tokens:,}"

    output_path = tempfile.mktemp(suffix=".docx")
    doc.save(output_path)
    return output_path


def _markdown_to_html(md: str) -> str:
    """Simple markdown-to-HTML conversion for PDF rendering."""
    import re

    html = md
    # Headers
    html = re.sub(r'^#### (.+)$', r'<h4>\1</h4>', html, flags=re.MULTILINE)
    html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
    html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
    html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)
    # Bold
    html = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', html)
    # Italic
    html = re.sub(r'\*([^*]+)\*', r'<em>\1</em>', html)
    # List items
    html = re.sub(r'^- (.+)$', r'<li>\1</li>', html, flags=re.MULTILINE)
    # Paragraphs
    html = re.sub(r'\n\n', '</p><p>', html)

    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
body {{ font-family: Georgia, serif; max-width: 800px; margin: 40px auto; line-height: 1.6; }}
h1, h2, h3 {{ color: #102030; }}
table {{ border-collapse: collapse; width: 100%; }}
th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
</style>
</head>
<body><p>{html}</p></body>
</html>"""
